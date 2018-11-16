# -*- coding: utf-8 -*-

from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH   # 用来居中显示标题
import sys
from time import sleep

print(sys.getdefaultencoding())

# 浏览器安装路径
BROWSER_PATH = \
    'C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe'
# 目的URL
DEST_URL = "https://wenku.baidu.com/view/aa31a84bcf84b9d528ea7a2c.html"

# 用来保存文档
doc_title = ''
doc_content_list = []


def find_doc(driver, init=True):
    '''滚动Browser滚动条，遍历所有文档，实现递归调用

    args:
        1. driver: 使用的驱动；
        2. init: True，说明是第一次调用，其搜索的元素的位置和后续的不同；
                 False，搜索时，使用的锁定位置的元素相同；

    注意事项：
        1. stop_condition
        使用 stop_condition 变量是因为需要判断什么时候停止。这里我才用的是寻找
        “点击继续阅读”这个元素来锁定滑动的位置。如果，没有找到这个元素（其类型为
        WebElement），说明已经到头，停止递归调用。
        2. driver.execute_script
        这个函数有多种方式滑动滚动条，这里使用arguments[0].scrollIntoView()，
        是指根据元素的位置，滑动到其位置。
        3. driver.find_element_by_xpath
        根据相对路径或者绝对路径查找元素。具体怎么使用可以自行百度。
        例如，
        driver.find_element_by_xpath("//div[@class='pagerwg-button']")
    '''
    global doc_content_list
    global doc_title
    stop_condition = False
    html = driver.page_source
    soup1 = BeautifulSoup(html, 'lxml')

    if (init is True):
        # 得到标题
        title_result = soup1.find('div', attrs={'class': 'doc-title'})
        doc_title = title_result.get_text()  # 得到文档标题

        # 拖动滚动条
        init_page = driver.find_element_by_xpath(
            "//div[@class='foldpagewg-text-con']")
        print(type(init_page), init_page)
        driver.execute_script('arguments[0].scrollIntoView();', init_page)
        init_page.click()
        init = False
    else:
        try:
            next_page = driver.find_element_by_xpath(
                "//div[@class='pagerwg-loadSucc']")

            print(type(next_page), next_page)
            driver.execute_script('arguments[0].scrollIntoView();', next_page)
            next_page = driver.find_element_by_xpath(
                "//div[@class='pagerwg-button']")
            next_page.click()
        except:
            stop_condition = True

    # 遍历所有的txt标签标定的文档，将其空格删除，然后进行保存
    content_result = soup1.find_all('p', attrs={'class': 'txt'})
    for each in content_result:
        each_text = each.get_text()
        if '            ' in each_text:
            text = each_text.replace('            ', '')
        else:
            text = each_text
        # print(each_text)
        doc_content_list.append(text)  # 得到正文内容

    sleep(2)  # 防止页面加载过慢
    if stop_condition is False:
        doc_title, doc_content_list = find_doc(driver, init)

    return doc_title, doc_content_list


def save(doc_title, doc_content_list):
    '''将获取的文档标题和文档内容写入到Word文档中
    args:
        1. doc_title: 文档标题，类型是字符串；
        2. doc_content_list： 文档内容，类型是列表；
    '''
    document = Document()
    heading = document.add_heading(doc_title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中显示

    for each in doc_content_list:
        document.add_paragraph(each)

    # 处理字符编码问题
    t_title = doc_title.encode('utf-8').split()[0]
    document.save('百度文库-%s.docx' % t_title)
    # document.save('2.docx')

    print("\n\nCompleted: %s.docx, to read." % t_title)
    driver.quit()


if __name__ == '__main__':
    options = webdriver.ChromeOptions()
    options.add_argument('user-agent="Mozilla/5.0 (Linux; Android 4.0.4; \
        Galaxy Nexus Build/IMM76B) AppleWebKit/535.19 (KHTML, like Gecko) \
        Chrome/18.0.1025.133 Mobile Safari/535.19"')
    driver = webdriver.Chrome(BROWSER_PATH, chrome_options=options)
    driver.get(DEST_URL)

    print("**********START**********")
    title, content = find_doc(driver, True)
    save(title, content)
